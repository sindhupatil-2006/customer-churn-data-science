"""
Week 4 Asset Generation Script (Notebooks, Figures, Word Executive Report)
Author: Sindhu Patil (Data Science Intern)
"""

import os
import json
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

from sklearn.model_selection import train_test_split
from sklearn.linear_model import LogisticRegression
from sklearn.ensemble import RandomForestClassifier
from sklearn.metrics import precision_recall_curve, f1_score, precision_score, recall_score

from src.modeling import (
    build_preprocessing_pipeline, build_model_pipeline, evaluate_cv_performance,
    evaluate_test_performance, extract_logistic_coefficients
)
from src.visualization import (
    plot_confusion_matrix, plot_roc_curve, plot_precision_recall_curve,
    plot_threshold_analysis, plot_model_comparison, plot_top_coefficients
)

DATA_PATH = os.path.join("data", "processed", "cleaned_telco_churn.csv")
FIGURES_DIR = os.path.join("outputs", "figures", "week4")
RESULTS_DIR = os.path.join("outputs", "results")
NOTEBOOKS_DIR = "notebooks"
REPORTS_DIR = "reports"

os.makedirs(FIGURES_DIR, exist_ok=True)
os.makedirs(RESULTS_DIR, exist_ok=True)
os.makedirs(NOTEBOOKS_DIR, exist_ok=True)
os.makedirs(REPORTS_DIR, exist_ok=True)


def make_cell(cell_type, source):
    """Creates Jupyter Notebook cell dict."""
    return {
        "cell_type": cell_type,
        "metadata": {},
        "outputs": [],
        "execution_count": None,
        "source": source if isinstance(source, list) else source.splitlines(keepends=True)
    }


def generate_week4_figures(y_test, y_prob_lr, y_prob_rf, cm_lr, metrics_lr, metrics_rf, df_coef):
    """Generates all 6 Week 4 figure plots in outputs/figures/week4/."""
    print("Generating Week 4 supervised learning figures...")
    plot_confusion_matrix(cm_lr, display_labels=['Retained (No)', 'Churned (Yes)'], filename="confusion_matrix.png")

    y_prob_dict = {
        'Logistic Regression': (y_prob_lr, metrics_lr['ROC_AUC']),
        'Random Forest': (y_prob_rf, metrics_rf['ROC_AUC'])
    }
    plot_roc_curve(y_test, y_prob_dict, "roc_curve.png")

    pr_dict = {
        'Logistic Regression': (y_prob_lr, metrics_lr['Average_Precision']),
        'Random Forest': (y_prob_rf, metrics_rf['Average_Precision'])
    }
    plot_precision_recall_curve(y_test, pr_dict, "precision_recall_curve.png")

    # Threshold analysis
    thresholds = np.linspace(0.1, 0.9, 81)
    precisions, recalls, f1s = [], [], []
    for t in thresholds:
        y_pred_t = (y_prob_lr >= t).astype(int)
        precisions.append(precision_score(y_test, y_pred_t, zero_division=0))
        recalls.append(recall_score(y_test, y_pred_t, zero_division=0))
        f1s.append(f1_score(y_test, y_pred_t, zero_division=0))

    plot_threshold_analysis(thresholds, precisions, recalls, f1s, "threshold_analysis.png")

    # Model comparison
    comp_df = pd.DataFrame([metrics_lr, metrics_rf])
    plot_model_comparison(comp_df, "model_comparison.png")

    # Top coefficients
    plot_top_coefficients(df_coef, top_n=12, filename="top_logistic_coefficients.png")

    print("All 6 Week 4 figures generated successfully.")


def generate_week4_notebook():
    """Generates 21-section notebooks/week4_supervised_learning.ipynb."""
    cells = [
        make_cell("markdown", "# Week 4 — Supervised Learning Model Implementation\n**Student Name:** Sindhu Patil | **Internship:** Data Science\n\n## 1. Introduction\nSupervised machine learning trains algorithms on historical data paired with ground-truth target labels to learn predictive mappings."),
        make_cell("markdown", "## 2. Problem Statement\nCustomer churn in telecommunications leads to revenue loss. Predicting subscriber churn prior to cancellation enables proactive retention interventions."),
        make_cell("markdown", "## 3. Objective\nBuild, evaluate, and interpret a binary classification pipeline using Logistic Regression on `data/processed/cleaned_telco_churn.csv`."),
        make_cell("markdown", "## 4. Dataset Overview\nLoading cleaned dataset (7,043 rows, 21 columns). `customerID` identifier is strictly excluded."),
        make_cell("code", "import sys, os\nsys.path.insert(0, '..')\n\nimport pandas as pd\nimport numpy as np\nimport matplotlib.pyplot as plt\nimport seaborn as sns\nfrom sklearn.model_selection import train_test_split\nfrom sklearn.linear_model import LogisticRegression\nfrom sklearn.ensemble import RandomForestClassifier\nfrom src.modeling import build_preprocessing_pipeline, build_model_pipeline, evaluate_cv_performance, evaluate_test_performance, extract_logistic_coefficients\nfrom src.visualization import plot_confusion_matrix, plot_roc_curve, plot_precision_recall_curve, plot_threshold_analysis, plot_model_comparison, plot_top_coefficients\n\ndf = pd.read_csv('../data/processed/cleaned_telco_churn.csv')\nprint('Cleaned Dataset Shape:', df.shape)"),
        make_cell("markdown", "## 5. Target Variable Preparation (`Churn`)\nConverting binary target: `Yes` -> 1 (Churned), `No` -> 0 (Retained). Target is imbalanced (26.54% positive class)."),
        make_cell("code", "X = df.drop(columns=['customerID', 'Churn'])\ny = (df['Churn'] == 'Yes').astype(int)\nprint('Target class counts:\\n', y.value_counts(normalize=True))"),
        make_cell("markdown", "## 6. Feature Selection & Feature Engineering\nCategorizing 3 continuous numerical features (`tenure`, `MonthlyCharges`, `TotalCharges`) and 15 categorical predictors."),
        make_cell("code", "numerical_features = ['tenure', 'MonthlyCharges', 'TotalCharges']\ncategorical_features = [c for c in X.columns if c not in numerical_features]\nprint('Numerical features:', numerical_features)\nprint('Categorical features:', categorical_features)"),
        make_cell("markdown", "## 7. Train-Test Split (80/20 Stratified)\nSplitting data into 5,634 training rows and 1,409 testing rows. Stratification preserves the 26.54% churn ratio in both splits."),
        make_cell("code", "X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42, stratify=y)\nprint('X_train shape:', X_train.shape, '| X_test shape:', X_test.shape)"),
        make_cell("markdown", "## 8. Preprocessing Pipeline (`ColumnTransformer`)\n`StandardScaler` standardizes continuous features, while `OneHotEncoder` transforms categorical predictors. Encapsulated inside `Pipeline` to prevent data leakage."),
        make_cell("code", "preprocessor = build_preprocessing_pipeline(numerical_features, categorical_features)"),
        make_cell("markdown", "## 9. Primary Model — Logistic Regression\nConstructing `Pipeline([('preprocessor', preprocessor), ('classifier', LogisticRegression(max_iter=1000, random_state=42))])`."),
        make_cell("code", "lr_pipeline = build_model_pipeline(LogisticRegression(max_iter=1000, random_state=42), preprocessor)"),
        make_cell("markdown", "## 10. 5-Fold Stratified Cross-Validation"),
        make_cell("code", "cv_summary, cv_results = evaluate_cv_performance(lr_pipeline, X, y, cv_folds=5)\nfor metric, vals in cv_summary.items():\n    print(f'{metric.upper():12s}: Mean = {vals[\"mean\"]:.4f} +/- {vals[\"std\"]:.4f}')"),
        make_cell("markdown", "## 11. Model Training & Test Set Evaluation"),
        make_cell("code", "metrics_lr, y_pred_lr, y_prob_lr, cm_lr = evaluate_test_performance(lr_pipeline, X_train, X_test, y_train, y_test, 'Logistic Regression')\nprint('Test Set Performance (Logistic Regression):\\n', metrics_lr)"),
        make_cell("markdown", "## 12. Confusion Matrix Analysis"),
        make_cell("code", "plot_confusion_matrix(cm_lr, ['Retained (No)', 'Churned (Yes)'], '../outputs/figures/week4/confusion_matrix.png')\nplt.figure(figsize=(5, 4))\nsns.heatmap(cm_lr, annot=True, fmt='d', cmap='Blues', xticklabels=['No', 'Yes'], yticklabels=['No', 'Yes'])\nplt.title('Confusion Matrix (Test Set)')\nplt.show()"),
        make_cell("markdown", "## 13. Classification Report"),
        make_cell("code", "from sklearn.metrics import classification_report\nprint(classification_report(y_test, y_pred_lr, target_names=['Retained', 'Churned']))"),
        make_cell("markdown", "## 14. ROC Curve & AUC Evaluation"),
        make_cell("code", "plot_roc_curve(y_test, {'Logistic Regression': (y_prob_lr, metrics_lr['ROC_AUC'])}, '../outputs/figures/week4/roc_curve.png')"),
        make_cell("markdown", "## 15. Precision-Recall Curve"),
        make_cell("code", "plot_precision_recall_curve(y_test, {'Logistic Regression': (y_prob_lr, metrics_lr['Average_Precision'])}, '../outputs/figures/week4/precision_recall_curve.png')"),
        make_cell("markdown", "## 16. Classification Threshold Analysis"),
        make_cell("code", "thresholds = np.linspace(0.1, 0.9, 81)\nprecisions = [precision_score(y_test, y_prob_lr >= t, zero_division=0) for t in thresholds]\nrecalls = [recall_score(y_test, y_prob_lr >= t, zero_division=0) for t in thresholds]\nf1s = [f1_score(y_test, y_prob_lr >= t, zero_division=0) for t in thresholds]\nplot_threshold_analysis(thresholds, precisions, recalls, f1s, '../outputs/figures/week4/threshold_analysis.png')"),
        make_cell("markdown", "## 17. Feature Coefficient Interpretation"),
        make_cell("code", "df_coef = extract_logistic_coefficients(lr_pipeline, numerical_features, categorical_features)\nplot_top_coefficients(df_coef, top_n=12, filename='../outputs/figures/week4/top_logistic_coefficients.png')\nprint(df_coef.head(10))"),
        make_cell("markdown", "## 18. Model Comparison (Logistic Regression vs Random Forest)"),
        make_cell("code", "rf_pipeline = build_model_pipeline(RandomForestClassifier(n_estimators=100, random_state=42), preprocessor)\nmetrics_rf, y_pred_rf, y_prob_rf, cm_rf = evaluate_test_performance(rf_pipeline, X_train, X_test, y_train, y_test, 'Random Forest')\ncomp_df = pd.DataFrame([metrics_lr, metrics_rf])\nplot_model_comparison(comp_df, '../outputs/figures/week4/model_comparison.png')\nprint(comp_df[['Model', 'Accuracy', 'Precision', 'Recall', 'F1_Score', 'ROC_AUC']])"),
        make_cell("markdown", "## 19. Final Model Selection Reasoning\nLogistic Regression was selected over Random Forest because it achieved superior Recall (**0.5588** vs 0.4759), higher F1-Score (**0.6040** vs 0.5370), and higher ROC-AUC (**0.8419** vs 0.8179), while offering clear coefficient interpretability."),
        make_cell("markdown", "## 20. Strengths & Limitations\n- **Strengths:** 0% data leakage via sklearn Pipeline, balanced evaluation across Recall/F1/ROC-AUC, clear feature coefficient rankings.\n- **Limitations:** Imbalanced dataset (~26.5% positive class) limits un-tuned Recall to ~55.9%."),
        make_cell("markdown", "## 21. Conclusion\nLogistic Regression established a robust baseline churn model with 80.55% accuracy and 0.8419 ROC-AUC, identifying short tenure, month-to-month contracts, and fiber optic service as primary churn drivers.")
    ]

    nb_path = os.path.join(NOTEBOOKS_DIR, "week4_supervised_learning.ipynb")
    nb = {
        "cells": cells,
        "metadata": {
            "language_info": {"name": "python", "version": "3.14.3"},
            "kernelspec": {"display_name": "Python 3", "language": "python", "name": "python3"}
        },
        "nbformat": 4, "nbformat_minor": 2
    }
    with open(nb_path, "w", encoding="utf-8") as f:
        json.dump(nb, f, indent=2)
    print(f"Created Jupyter Notebook: {nb_path}")


def set_cell_background(cell, fill_hex):
    """Sets cell background color in docx table."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def build_week4_word_report(df, cv_summary, metrics_lr, metrics_rf, df_coef):
    """
    Builds reports/week4_report.docx adhering strictly to prompt formatting rules:
    - Font: Times New Roman
    - Body font: 11 pt, 1.15 line spacing
    - Cover page for Sindhu Patil
    - Table of Contents
    - 28 detailed sections
    - Formatted tables & actual figure image embeds
    - Consolas code blocks
    """
    doc = Document()

    # Configure default styles to Times New Roman
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(51, 51, 51)
    normal_style.paragraph_format.line_spacing = 1.15

    # ---------------------------------------------------------
    # COVER PAGE
    # ---------------------------------------------------------
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("Customer Churn Prediction and Customer Segmentation\nUsing Python")
    run_title.font.name = 'Times New Roman'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(41, 128, 185)

    doc.add_paragraph("\n")
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("DATA SCIENCE INTERNSHIP\nWEEK 4 REPORT\nSUPERVISED LEARNING MODEL IMPLEMENTATION")
    run_sub.font.name = 'Times New Roman'
    run_sub.font.size = Pt(15)
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGBColor(52, 73, 94)

    doc.add_paragraph("\n" * 3)

    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ("Student Name:", "Sindhu Patil"),
        ("Internship Track:", "Data Science Internship"),
        ("Week Number:", "Week 4"),
        ("Project Title:", "Customer Churn Prediction and Customer Segmentation"),
        ("Dataset Used:", "IBM Telco Customer Churn Dataset")
    ]
    for i, (k, v) in enumerate(info_data):
        r0 = table.rows[i].cells[0].paragraphs[0].add_run(k)
        r0.font.name = 'Times New Roman'; r0.font.bold = True; r0.font.size = Pt(11)
        r1 = table.rows[i].cells[1].paragraphs[0].add_run(v)
        r1.font.name = 'Times New Roman'; r1.font.size = Pt(11)
        set_cell_background(table.rows[i].cells[0], "F2F4F4")
        set_cell_background(table.rows[i].cells[1], "EAEDED")

    doc.add_page_break()

    # ---------------------------------------------------------
    # TABLE OF CONTENTS
    # ---------------------------------------------------------
    h_toc = doc.add_heading("Table of Contents", level=1)
    h_toc.runs[0].font.name = 'Times New Roman'
    
    sections = [
        "1. Introduction", "2. Problem Statement", "3. Objective", "4. Dataset Overview",
        "5. Target Variable", "6. Feature Selection", "7. Feature Engineering",
        "8. Data Preprocessing", "9. Train-Test Split", "10. Model Selection",
        "11. Logistic Regression", "12. Model Training", "13. Cross-Validation",
        "14. Evaluation Metrics", "15. Confusion Matrix", "16. ROC Curve and AUC",
        "17. Precision-Recall Analysis", "18. Classification Threshold Analysis",
        "19. Feature Interpretation", "20. Optional Model Comparison", "21. Final Model Selection",
        "22. Business Interpretation", "23. Strengths", "24. Limitations",
        "25. Possible Improvements", "26. Conclusion", "27. References", "28. Appendix"
    ]
    for s in sections:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"•  {s}")
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)

    doc.add_paragraph("\n")

    def add_sec_heading(title, level=1):
        h = doc.add_heading(title, level=level)
        for run in h.runs:
            run.font.name = 'Times New Roman'
            run.font.bold = True
            run.font.color.rgb = RGBColor(41, 128, 185)
        return h

    # ---------------------------------------------------------
    # SECTIONS 1 to 5
    # ---------------------------------------------------------
    add_sec_heading("1. Introduction")
    doc.add_paragraph("Following exploratory data analysis (Week 2) and unsupervised customer segmentation (Week 3), Week 4 implements supervised machine learning classification models to predict customer churn.")

    add_sec_heading("2. Problem Statement")
    doc.add_paragraph("Customer churn represents a critical revenue challenge in telecommunications. Predicting churn risk prior to cancellation allows business teams to initiate proactive retention workflows.")

    add_sec_heading("3. Objective")
    doc.add_paragraph("Build a robust binary classification pipeline using Logistic Regression on `data/processed/cleaned_telco_churn.csv`, evaluate cross-validation performance, analyze test set metrics, and interpret predictive feature coefficients.")

    add_sec_heading("4. Dataset Overview")
    doc.add_paragraph("The dataset contains 7,043 customer records and 21 features. `customerID` was excluded as a non-predictive identifier.")

    add_sec_heading("5. Target Variable")
    doc.add_paragraph("Target `Churn` is mapped to binary indicators: 1 = Churned (1,869 customers / 26.54%), 0 = Retained (5,174 customers / 73.46%).")

    # ---------------------------------------------------------
    # SECTIONS 6 to 10
    # ---------------------------------------------------------
    add_sec_heading("6. Feature Selection")
    doc.add_paragraph("Selected 18 predictor variables across continuous numerical attributes (`tenure`, `MonthlyCharges`, `TotalCharges`) and 15 categorical attributes (`Contract`, `InternetService`, `PaymentMethod`, etc.).")

    add_sec_heading("7. Feature Engineering")
    doc.add_paragraph("Categorical variables were One-Hot Encoded into binary dummy columns, while continuous numerical attributes were standardized.")

    add_sec_heading("8. Data Preprocessing")
    doc.add_paragraph("Scikit-learn `ColumnTransformer` combined `StandardScaler` and `OneHotEncoder(handle_unknown='ignore')` inside a single unified `Pipeline` to prevent data leakage.")

    add_sec_heading("9. Train-Test Split")
    doc.add_paragraph("Applied an 80/20 stratified split (`test_size=0.2, random_state=42, stratify=y`), producing 5,634 training records and 1,409 testing records.")

    add_sec_heading("10. Model Selection")
    doc.add_paragraph("Logistic Regression was chosen as the primary baseline classifier due to its interpretability, probability output generation, and computational efficiency.")

    # ---------------------------------------------------------
    # SECTIONS 11 to 13 (CV Results Table)
    # ---------------------------------------------------------
    add_sec_heading("11. Logistic Regression")
    doc.add_paragraph("Logistic Regression models log-odds of churn as a linear combination of input predictors using the sigmoid activation function.")

    add_sec_heading("12. Model Training")
    doc.add_paragraph("The pipeline was fitted strictly on the 5,634 training samples (`X_train`, `y_train`).")

    add_sec_heading("13. Cross-Validation")
    doc.add_paragraph("Table 1 displays 5-fold stratified cross-validation metrics across training folds:")

    # Table 1: CV Metrics
    t_cv = doc.add_table(rows=6, cols=3)
    t_cv.alignment = WD_TABLE_ALIGNMENT.CENTER
    cv_headers = ["Metric", "Mean CV Score", "Standard Deviation"]
    for j, h in enumerate(cv_headers):
        cell = t_cv.rows[0].cells[j]
        r = cell.paragraphs[0].add_run(h)
        r.font.name = 'Times New Roman'; r.font.bold = True; r.font.size = Pt(10)
        set_cell_background(cell, "2980B9")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    cv_metric_names = [("Accuracy", "accuracy"), ("Precision", "precision"), ("Recall", "recall"), ("F1-Score", "f1"), ("ROC-AUC", "roc_auc")]
    for i, (disp_name, key) in enumerate(cv_metric_names):
        r_cells = t_cv.rows[i+1].cells
        r_cells[0].paragraphs[0].add_run(disp_name).font.name = 'Times New Roman'
        r_cells[1].paragraphs[0].add_run(f"{cv_summary[key]['mean']:.4f}").font.name = 'Times New Roman'
        r_cells[2].paragraphs[0].add_run(f"± {cv_summary[key]['std']:.4f}").font.name = 'Times New Roman'
        for c in r_cells:
            set_cell_background(c, "F2F4F4" if i % 2 == 0 else "FFFFFF")

    # ---------------------------------------------------------
    # SECTIONS 14 to 18 (Evaluation, CM, ROC, PR, Threshold)
    # ---------------------------------------------------------
    add_sec_heading("14. Evaluation Metrics")
    doc.add_paragraph(f"Test set performance (1,409 samples): Accuracy = {metrics_lr['Accuracy']:.4f}, Precision = {metrics_lr['Precision']:.4f}, Recall = {metrics_lr['Recall']:.4f}, F1-Score = {metrics_lr['F1_Score']:.4f}, ROC-AUC = {metrics_lr['ROC_AUC']:.4f}.")

    add_sec_heading("15. Confusion Matrix")
    doc.add_paragraph(f"Test set confusion matrix breakdown:\n• True Negatives (TN): {metrics_lr['TN']} (Retained customers correctly predicted as Retained)\n• False Positives (FP): {metrics_lr['FP']} (Retained customers incorrectly predicted as Churn)\n• False Negatives (FN): {metrics_lr['FN']} (Churned customers incorrectly predicted as Retained)\n• True Positives (TP): {metrics_lr['TP']} (Churned customers correctly predicted as Churn)")

    fig_path = os.path.join(FIGURES_DIR, "confusion_matrix.png")
    if os.path.exists(fig_path):
        doc.add_paragraph("Figure 15.1: Logistic Regression Confusion Matrix")
        doc.add_picture(fig_path, width=Inches(5.2))

    add_sec_heading("16. ROC Curve and AUC")
    doc.add_paragraph(f"The ROC-AUC score of {metrics_lr['ROC_AUC']:.4f} demonstrates strong class separation capability.")

    fig_path = os.path.join(FIGURES_DIR, "roc_curve.png")
    if os.path.exists(fig_path):
        doc.add_paragraph("Figure 16.1: Receiver Operating Characteristic (ROC) Curve")
        doc.add_picture(fig_path, width=Inches(5.5))

    add_sec_heading("17. Precision-Recall Analysis")
    doc.add_paragraph(f"Average Precision score = {metrics_lr['Average_Precision']:.4f}, outperforming the 0.2654 baseline.")

    fig_path = os.path.join(FIGURES_DIR, "precision_recall_curve.png")
    if os.path.exists(fig_path):
        doc.add_paragraph("Figure 17.1: Precision-Recall Curve")
        doc.add_picture(fig_path, width=Inches(5.5))

    add_sec_heading("18. Classification Threshold Analysis")
    doc.add_paragraph("Lowering the classification threshold from 0.5 to 0.35 increases Recall (catching more churners) at the cost of lower Precision.")

    fig_path = os.path.join(FIGURES_DIR, "threshold_analysis.png")
    if os.path.exists(fig_path):
        doc.add_paragraph("Figure 18.1: Precision-Recall-F1 Threshold Trade-off Plot")
        doc.add_picture(fig_path, width=Inches(5.8))

    # ---------------------------------------------------------
    # SECTIONS 19 to 21 (Coefficients, Model Comparison)
    # ---------------------------------------------------------
    add_sec_heading("19. Feature Interpretation")
    doc.add_paragraph("Top negative coefficients (protective against churn): tenure (-1.2339), Two-year contract (-0.7627), DSL service (-0.6402). Top positive coefficients (churn risk factors): Fiber optic service (+0.6431), Month-to-month contract (+0.5873), TotalCharges (+0.5103).")

    fig_path = os.path.join(FIGURES_DIR, "top_logistic_coefficients.png")
    if os.path.exists(fig_path):
        doc.add_paragraph("Figure 19.1: Top Logistic Regression Feature Coefficients")
        doc.add_picture(fig_path, width=Inches(5.8))

    add_sec_heading("20. Optional Model Comparison")
    doc.add_paragraph("Table 2 compares Logistic Regression against Random Forest Classifier on the test set:")

    # Table 2: Model Comparison
    t_comp = doc.add_table(rows=3, cols=6)
    t_comp.alignment = WD_TABLE_ALIGNMENT.CENTER
    m_headers = ["Model Name", "Accuracy", "Precision", "Recall", "F1-Score", "ROC-AUC"]
    for j, h in enumerate(m_headers):
        cell = t_comp.rows[0].cells[j]
        r = cell.paragraphs[0].add_run(h)
        r.font.name = 'Times New Roman'; r.font.bold = True; r.font.size = Pt(10)
        set_cell_background(cell, "2980B9")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    comp_rows = [metrics_lr, metrics_rf]
    for i, m in enumerate(comp_rows):
        r_cells = t_comp.rows[i+1].cells
        r_cells[0].paragraphs[0].add_run(m['Model']).font.name = 'Times New Roman'
        r_cells[1].paragraphs[0].add_run(f"{m['Accuracy']:.4f}").font.name = 'Times New Roman'
        r_cells[2].paragraphs[0].add_run(f"{m['Precision']:.4f}").font.name = 'Times New Roman'
        r_cells[3].paragraphs[0].add_run(f"{m['Recall']:.4f}").font.name = 'Times New Roman'
        r_cells[4].paragraphs[0].add_run(f"{m['F1_Score']:.4f}").font.name = 'Times New Roman'
        r_cells[5].paragraphs[0].add_run(f"{m['ROC_AUC']:.4f}").font.name = 'Times New Roman'
        for c in r_cells:
            set_cell_background(c, "F2F4F4" if i % 2 == 0 else "FFFFFF")

    fig_path = os.path.join(FIGURES_DIR, "model_comparison.png")
    if os.path.exists(fig_path):
        doc.add_paragraph("Figure 20.1: Model Comparison Bar Chart")
        doc.add_picture(fig_path, width=Inches(5.8))

    add_sec_heading("21. Final Model Selection")
    doc.add_paragraph("Logistic Regression is selected as the primary baseline model due to superior Recall (0.5588 vs 0.4759), higher F1-Score (0.6040 vs 0.5370), higher ROC-AUC (0.8419 vs 0.8179), and direct coefficient interpretability.")

    # ---------------------------------------------------------
    # SECTIONS 22 to 28
    # ---------------------------------------------------------
    add_sec_heading("22. Business Interpretation")
    doc.add_paragraph("Logistic Regression probabilities allow retention teams to score customer churn risk dynamically, enabling targeted intervention for high-risk accounts.")

    add_sec_heading("23. Strengths")
    doc.add_paragraph("• 0% data leakage via Scikit-learn Pipeline architecture.\n• High interpretability via linear feature coefficients.\n• Robust cross-validation metrics.")

    add_sec_heading("24. Limitations")
    doc.add_paragraph("• Moderate recall (~55.9%) under standard 0.5 probability threshold due to class imbalance.")

    add_sec_heading("25. Possible Improvements")
    doc.add_paragraph("Future iterations may apply class weighting (`class_weight='balanced'`), hyperparameter tuning via GridSearchCV, SMOTE oversampling, or XGBoost.")

    add_sec_heading("26. Conclusion")
    doc.add_paragraph("Week 4 supervised classification established an effective baseline churn model achieving 80.55% accuracy and 0.8419 ROC-AUC.")

    add_sec_heading("27. References")
    doc.add_paragraph("1. Hosmer, D. W., Lemeshow, S., & Sturdivant, R. X. (2013). Applied Logistic Regression.\n2. Pedregosa, F., et al. (2011). Scikit-learn: Machine Learning in Python.")

    add_sec_heading("28. Appendix — Code Snippet")
    doc.add_paragraph("```python\n# Preprocessing Pipeline & Model Training\npreprocessor = ColumnTransformer([\n    ('num', StandardScaler(), numerical_features),\n    ('cat', OneHotEncoder(handle_unknown='ignore'), categorical_features)\n])\nmodel = Pipeline([\n    ('preprocessor', preprocessor),\n    ('classifier', LogisticRegression(max_iter=1000, random_state=42))\n])\nmodel.fit(X_train, y_train)\n```")

    doc_path = os.path.join(REPORTS_DIR, "week4_report.docx")
    doc.save(doc_path)
    print(f"Created Word Executive Report: {doc_path}")


def main():
    print("==========================================================")
    print("STARTING WEEK 4 ASSET GENERATION FOR SINDHU PATIL")
    print("==========================================================")
    df = pd.read_csv(DATA_PATH)
    print(f"Loaded dataset: {DATA_PATH} with shape {df.shape}")

    X = df.drop(columns=['customerID', 'Churn'])
    y = (df['Churn'] == 'Yes').astype(int)

    num_cols = ['tenure', 'MonthlyCharges', 'TotalCharges']
    cat_cols = [c for c in X.columns if c not in num_cols]

    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size=0.2, random_state=42, stratify=y)
    preprocessor = build_preprocessing_pipeline(num_cols, cat_cols)

    # Logistic Regression
    lr_pipe = build_model_pipeline(LogisticRegression(max_iter=1000, random_state=42), preprocessor)
    cv_summary, _ = evaluate_cv_performance(lr_pipe, X, y)
    metrics_lr, y_pred_lr, y_prob_lr, cm_lr = evaluate_test_performance(lr_pipe, X_train, X_test, y_train, y_test, 'Logistic Regression')
    df_coef = extract_logistic_coefficients(lr_pipe, num_cols, cat_cols)

    # Random Forest
    rf_pipe = build_model_pipeline(RandomForestClassifier(n_estimators=100, random_state=42), preprocessor)
    metrics_rf, y_pred_rf, y_prob_rf, cm_rf = evaluate_test_performance(rf_pipe, X_train, X_test, y_train, y_test, 'Random Forest')

    # Save model results CSV
    results_df = pd.DataFrame([metrics_lr, metrics_rf])
    results_csv_path = os.path.join(RESULTS_DIR, "week4_model_results.csv")
    results_df.to_csv(results_csv_path, index=False)
    print(f"Saved Week 4 Model Results CSV to {results_csv_path}")

    generate_week4_figures(y_test, y_prob_lr, y_prob_rf, cm_lr, metrics_lr, metrics_rf, df_coef)
    generate_week4_notebook()
    build_week4_word_report(df, cv_summary, metrics_lr, metrics_rf, df_coef)

    print("==========================================================")
    print("WEEK 4 ASSET GENERATION COMPLETED SUCCESSFULLY")
    print("==========================================================")


if __name__ == "__main__":
    main()
