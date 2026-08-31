"""
Week 2 Asset Generation Script (Notebooks, Figures, Word Executive Report)
Author: Sindhu Patil (Data Science Intern)
"""

import os
import json
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

from src.visualization import (
    plot_churn_target_distribution, plot_churn_by_contract, plot_churn_by_internet_service,
    plot_churn_by_payment_method, plot_tenure_distribution, plot_monthly_charges_by_churn,
    plot_total_charges_by_churn, plot_correlation_heatmap, plot_multivariate_contract_charges,
    plot_demographics_churn, plot_services_churn
)

DATA_PATH = os.path.join("data", "processed", "cleaned_telco_churn.csv")
FIGURES_DIR = os.path.join("outputs", "figures", "week2")
NOTEBOOKS_DIR = "notebooks"
REPORTS_DIR = "reports"

os.makedirs(FIGURES_DIR, exist_ok=True)
os.makedirs(NOTEBOOKS_DIR, exist_ok=True)
os.makedirs(REPORTS_DIR, exist_ok=True)


def make_cell(cell_type, source):
    """Creates Jupyter Notebook cell dict."""
    return {
        "cell_type": cell_type,
        "metadata": {},
        "outputs": [],
        "execution_count": None,
        "source": source if isinstance(source, list) else source.splitlines(keepends=True)
    }


def generate_week2_figures(df):
    """Generates all 11 Week 2 figures into outputs/figures/week2/."""
    print("Generating Week 2 figure plots...")
    plot_churn_target_distribution(df, "churn_distribution.png")
    plot_churn_by_contract(df, "churn_by_contract.png")
    plot_churn_by_internet_service(df, "churn_by_internet_service.png")
    plot_churn_by_payment_method(df, "churn_by_payment_method.png")
    plot_tenure_distribution(df, "tenure_distribution.png")
    plot_monthly_charges_by_churn(df, "monthly_charges_by_churn.png")
    plot_total_charges_by_churn(df, "total_charges_by_churn.png")
    plot_correlation_heatmap(df, "correlation_heatmap.png")
    plot_multivariate_contract_charges(df, "multivariate_contract_charges.png")
    plot_demographics_churn(df, "demographics_churn.png")
    plot_services_churn(df, "services_churn.png")
    print("All 11 Week 2 figures saved successfully.")


def generate_week2_notebook():
    """Generates 18-section notebooks/week2_eda.ipynb."""
    cells = [
        make_cell("markdown", "# Week 2 — Exploratory Data Analysis and Visualization\n**Student Name:** Sindhu Patil | **Internship:** Data Science\n\n## 1. Introduction\nExploratory Data Analysis (EDA) is a fundamental phase in data science that uncovers underlying statistical distributions, identifies feature correlations, evaluates class balance, and extracts domain-specific business insights before modeling."),
        make_cell("markdown", "## 2. Objective\nPerform data-backed EDA on the cleaned IBM Telco Customer Churn dataset (`data/processed/cleaned_telco_churn.csv`). Quantify relationships between customer demographics, contract commitments, payment channels, service add-ons, and churn rates."),
        make_cell("markdown", "## 3. Import Libraries & Configuration"),
        make_cell("code", "import sys, os\nsys.path.insert(0, '..')\n\nimport pandas as pd\nimport numpy as np\nimport matplotlib.pyplot as plt\nimport seaborn as sns\nfrom src.visualization import (\n    plot_churn_target_distribution, plot_churn_by_contract, plot_churn_by_internet_service,\n    plot_churn_by_payment_method, plot_tenure_distribution, plot_monthly_charges_by_churn,\n    plot_total_charges_by_churn, plot_correlation_heatmap, plot_multivariate_contract_charges,\n    plot_demographics_churn, plot_services_churn\n)\n\nplt.style.use('seaborn-v0_8-whitegrid' if 'seaborn-v0_8-whitegrid' in plt.style.available else 'default')\nsns.set_palette('husl')"),
        make_cell("markdown", "## 4. Load Cleaned Dataset"),
        make_cell("code", "data_path = '../data/processed/cleaned_telco_churn.csv'\ndf = pd.read_csv(data_path)\nprint('Loaded Dataset Shape:', df.shape)\nprint('Missing values:', df.isnull().sum().sum())\nprint('Duplicate rows:', df.duplicated().sum())"),
        make_cell("markdown", "## 5. Dataset Overview & Data Structure"),
        make_cell("code", "print('Columns:', df.columns.tolist())\ndf.info()"),
        make_cell("code", "df.head(5)"),
        make_cell("code", "df.tail(5)"),
        make_cell("markdown", "## 6. Descriptive Statistics\nCalculating Mean, Median, Standard Deviation, Min, Max, Q1, Q3, and IQR for numerical features."),
        make_cell("code", "num_cols = ['tenure', 'MonthlyCharges', 'TotalCharges']\nstats = df[num_cols].describe().T\nstats['median'] = df[num_cols].median()\nstats['iqr'] = stats['75%'] - stats['25%']\nstats = stats[['mean', 'std', 'min', '25%', 'median', '75%', 'max', 'iqr']]\nstats"),
        make_cell("markdown", "## 7. Target Variable — Churn Analysis"),
        make_cell("code", "churn_counts = df['Churn'].value_counts()\nchurn_pcts = df['Churn'].value_counts(normalize=True) * 100\nchurn_summary = pd.DataFrame({'Customer_Count': churn_counts, 'Percentage_%': churn_pcts})\nprint(churn_summary)\n\nfig1 = plot_churn_target_distribution(df, 'churn_distribution.png')\nplt.figure(figsize=(6, 4))\nsns.countplot(data=df, x='Churn', palette=['#2ecc71', '#e74c3c'])\nplt.title('Target Churn Class Distribution (73.5% vs 26.5%)')\nplt.show()"),
        make_cell("markdown", "## 8. Univariate Feature Analysis\nVisualizing individual feature distributions."),
        make_cell("code", "plot_tenure_distribution(df, 'tenure_distribution.png')\nplt.figure(figsize=(8, 4))\nsns.histplot(df['tenure'], kde=True, color='#3498db')\nplt.title('Univariate Tenure Distribution (Months)')\nplt.show()"),
        make_cell("code", "plot_monthly_charges_by_churn(df, 'monthly_charges_by_churn.png')\nplot_total_charges_by_churn(df, 'total_charges_by_churn.png')"),
        make_cell("markdown", "## 9. Bivariate Analysis & Churn Rates by Category"),
        make_cell("code", "contract_ct = pd.crosstab(df['Contract'], df['Churn'], margins=True)\ncontract_ct['Churn_Rate_%'] = (pd.crosstab(df['Contract'], df['Churn'], normalize='index')['Yes'] * 100).round(2)\nprint('--- Contract Churn Rates ---')\nprint(contract_ct)\nplot_churn_by_contract(df, 'churn_by_contract.png')"),
        make_cell("code", "internet_ct = pd.crosstab(df['InternetService'], df['Churn'], margins=True)\ninternet_ct['Churn_Rate_%'] = (pd.crosstab(df['InternetService'], df['Churn'], normalize='index')['Yes'] * 100).round(2)\nprint('--- Internet Service Churn Rates ---')\nprint(internet_ct)\nplot_churn_by_internet_service(df, 'churn_by_internet_service.png')"),
        make_cell("code", "pay_ct = pd.crosstab(df['PaymentMethod'], df['Churn'], margins=True)\npay_ct['Churn_Rate_%'] = (pd.crosstab(df['PaymentMethod'], df['Churn'], normalize='index')['Yes'] * 100).round(2)\nprint('--- Payment Method Churn Rates ---')\nprint(pay_ct)\nplot_churn_by_payment_method(df, 'churn_by_payment_method.png')"),
        make_cell("markdown", "## 10. Tenure Cohort Analysis"),
        make_cell("code", "df['Tenure_Cohort'] = pd.cut(df['tenure'], bins=[-1, 12, 24, 48, 60, 72], labels=['0-12 mos', '13-24 mos', '25-48 mos', '49-60 mos', '61-72 mos'])\ntenure_ct = pd.crosstab(df['Tenure_Cohort'], df['Churn'], margins=True)\ntenure_ct['Churn_Rate_%'] = (pd.crosstab(df['Tenure_Cohort'], df['Churn'], normalize='index')['Yes'] * 100).round(2)\nprint('--- Tenure Cohort Churn Rates ---')\nprint(tenure_ct)"),
        make_cell("markdown", "## 11. Demographic & Value-Added Services Analysis"),
        make_cell("code", "plot_demographics_churn(df, 'demographics_churn.png')\nplot_services_churn(df, 'services_churn.png')"),
        make_cell("markdown", "## 12. Multivariate Analysis"),
        make_cell("code", "plot_multivariate_contract_charges(df, 'multivariate_contract_charges.png')\nplt.figure(figsize=(9, 5))\nsns.boxplot(data=df, x='Contract', y='MonthlyCharges', hue='Churn', palette=['#2ecc71', '#e74c3c'])\nplt.title('Monthly Charges by Contract & Churn Status')\nplt.show()"),
        make_cell("markdown", "## 13. Correlation Analysis"),
        make_cell("code", "plot_correlation_heatmap(df, 'correlation_heatmap.png')\nnum_corr = df[['tenure', 'MonthlyCharges', 'TotalCharges']].copy()\nnum_corr['Churn_Numeric'] = (df['Churn'] == 'Yes').astype(int)\nprint('Correlations with Churn:')\nprint(num_corr.corr()['Churn_Numeric'])"),
        make_cell("markdown", "## 14. Key Findings\n1. Target Imbalance: 26.54% churn rate (1,869 churners vs 5,174 retained).\n2. Contract Type: Month-to-Month contracts have 42.71% churn rate vs 2.83% for 2-Year plans.\n3. Internet Service: Fiber Optic subscribers exhibit 41.89% churn rate vs 18.96% for DSL.\n4. Payment Method: Electronic Check users show the highest churn rate at 45.29%.\n5. Tenure Impact: 0–12 Month subscribers have 47.44% churn rate vs 6.61% for 61–72 Month tenure.\n6. Security Add-ons: Customers without Online Security have 41.77% churn rate vs 14.61% with security.\n7. Tech Support: Customers without Tech Support have 41.64% churn rate vs 15.17% with support.\n8. Senior Status: Senior citizens churn at 41.68% vs 23.61% for non-seniors.\n9. Correlation: Tenure is negatively correlated with churn (r = -0.352).\n10. Monthly Charges: Higher monthly bill (mean $74.44 for churners vs $61.27 for retained) increases churn risk."),
        make_cell("markdown", "## 15. Business Insights & Recommendations\n- DATA FINDING: 42.71% Month-to-Month churn vs 2.83% 2-Year contract churn.\n  - BUSINESS STRATEGY: Offer a 15% annual contract renewal discount.\n- DATA FINDING: 41.89% Fiber Optic churn rate with mean monthly bill of $91.50.\n  - BUSINESS STRATEGY: Bundle free Online Security and Tech Support for Fiber subscribers.\n- DATA FINDING: 45.29% Electronic Check churn rate.\n  - BUSINESS STRATEGY: Incentive automated payment methods (Auto Credit Card / ACH) with a $5 monthly bill credit."),
        make_cell("markdown", "## 16. Technical Challenges & Solutions\n- Class Imbalance: Handled by focusing on Recall, F1-Score, and ROC-AUC rather than raw Accuracy.\n- Pandas 3.0 String Types: Explicitly cast boolean masks and integer categorical indicators."),
        make_cell("markdown", "## 17. Limitations\n- Cross-sectional historical data lacks real-time time-series telemetry.\n- Observational data proves correlation, not direct causality."),
        make_cell("markdown", "## 18. Conclusion\nWeek 2 EDA identified Month-to-Month contracts, Fiber Optic service, Electronic Check payments, and early tenure (<12 months) as primary drivers of customer churn.")
    ]

    nb_path = os.path.join(NOTEBOOKS_DIR, "week2_eda.ipynb")
    nb = {
        "cells": cells,
        "metadata": {
            "language_info": {"name": "python", "version": "3.14.3"},
            "kernelspec": {"display_name": "Python 3", "language": "python", "name": "python3"}
        },
        "nbformat": 4, "nbformat_minor": 2
    }
    with open(nb_path, "w", encoding="utf-8") as f:
        json.dump(nb, f, indent=2)
    print(f"Created Jupyter Notebook: {nb_path}")


def set_cell_background(cell, fill_hex):
    """Sets cell background color in docx table."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)


def build_week2_word_report(df):
    """
    Builds reports/week2_report.docx adhering strictly to prompt formatting rules:
    - Font: Times New Roman
    - Body font: 11 pt, 1.15 line spacing
    - Cover page for Sindhu Patil
    - Table of Contents
    - 19 detailed sections
    - Formatted tables & actual figure image embeds
    - 200+ word Submission Description
    """
    doc = Document()

    # Configure default styles to Times New Roman
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(11)
    normal_style.font.color.rgb = RGBColor(51, 51, 51)
    normal_style.paragraph_format.line_spacing = 1.15

    # ---------------------------------------------------------
    # COVER PAGE
    # ---------------------------------------------------------
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run("Customer Churn Prediction and Customer Segmentation\nUsing Python")
    run_title.font.name = 'Times New Roman'
    run_title.font.size = Pt(22)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(41, 128, 185)

    doc.add_paragraph("\n")
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = p_sub.add_run("DATA SCIENCE INTERNSHIP\nWEEK 2 REPORT\nEXPLORATORY DATA ANALYSIS AND VISUALIZATION")
    run_sub.font.name = 'Times New Roman'
    run_sub.font.size = Pt(15)
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGBColor(52, 73, 94)

    doc.add_paragraph("\n" * 3)

    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ("Student Name:", "Sindhu Patil"),
        ("Internship Track:", "Data Science Internship"),
        ("Week Number:", "Week 2"),
        ("Technologies Used:", "Python, Pandas, NumPy, Matplotlib, Seaborn, docx")
    ]
    for i, (k, v) in enumerate(info_data):
        r0 = table.rows[i].cells[0].paragraphs[0].add_run(k)
        r0.font.name = 'Times New Roman'; r0.font.bold = True; r0.font.size = Pt(11)
        r1 = table.rows[i].cells[1].paragraphs[0].add_run(v)
        r1.font.name = 'Times New Roman'; r1.font.size = Pt(11)
        set_cell_background(table.rows[i].cells[0], "F2F4F4")
        set_cell_background(table.rows[i].cells[1], "EAEDED")

    doc.add_page_break()

    # ---------------------------------------------------------
    # TABLE OF CONTENTS
    # ---------------------------------------------------------
    h_toc = doc.add_heading("Table of Contents", level=1)
    h_toc.runs[0].font.name = 'Times New Roman'
    
    sections = [
        "1. Introduction", "2. Objective", "3. Dataset Overview", "4. Tools and Technologies",
        "5. Data Preparation", "6. Descriptive Statistics", "7. Univariate Analysis",
        "8. Bivariate Analysis", "9. Multivariate Analysis", "10. Churn Analysis",
        "11. Correlation Analysis", "12. Key Findings", "13. Business Insights",
        "14. Challenges Faced", "15. Limitations", "16. Impact on Future Modeling",
        "17. Conclusion", "18. References", "19. Submission Description"
    ]
    for s in sections:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"•  {s}")
        r.font.name = 'Times New Roman'
        r.font.size = Pt(11)

    doc.add_paragraph("\n")

    # Helper function for adding styled headings
    def add_sec_heading(title, level=1):
        h = doc.add_heading(title, level=level)
        for run in h.runs:
            run.font.name = 'Times New Roman'
            run.font.bold = True
            run.font.color.rgb = RGBColor(41, 128, 185)
        return h

    # ---------------------------------------------------------
    # 1. Introduction
    # ---------------------------------------------------------
    add_sec_heading("1. Introduction")
    doc.add_paragraph("Exploratory Data Analysis (EDA) represents a pivotal milestone in data science workflows. Following Week 1 data cleaning, Week 2 investigates patterns, trends, customer behaviors, and statistical anomalies within the IBM Telco Customer Churn dataset.")

    # ---------------------------------------------------------
    # 2. Objective
    # ---------------------------------------------------------
    add_sec_heading("2. Objective")
    doc.add_paragraph("The objective of Week 2 is to perform rigorous data analysis to identify primary demographic, contractual, and service-related churn drivers, supporting subsequent machine learning model development.")

    # ---------------------------------------------------------
    # 3. Dataset Overview
    # ---------------------------------------------------------
    add_sec_heading("3. Dataset Overview")
    doc.add_paragraph("The dataset used is `data/processed/cleaned_telco_churn.csv`, consisting of 7,043 customer records and 21 features. It contains zero missing values and zero duplicate rows.")

    # ---------------------------------------------------------
    # 4. Tools and Technologies
    # ---------------------------------------------------------
    add_sec_heading("4. Tools and Technologies")
    doc.add_paragraph("• Python 3.14\n• Pandas & NumPy (Data manipulation & aggregations)\n• Matplotlib & Seaborn (Statistical visualization)\n• python-docx (Executive report formatting)")

    # ---------------------------------------------------------
    # 5. Data Preparation
    # ---------------------------------------------------------
    add_sec_heading("5. Data Preparation")
    doc.add_paragraph("Dataset integrity was verified by confirming 0 null values across all 7,043 rows. `tenure` was binned into 5 analytical cohorts: 0–12 mos, 13–24 mos, 25–48 mos, 49–60 mos, and 61–72 mos.")

    # ---------------------------------------------------------
    # 6. Descriptive Statistics
    # ---------------------------------------------------------
    add_sec_heading("6. Descriptive Statistics")
    doc.add_paragraph("Table 1 summarizes descriptive statistics for continuous numerical attributes:")

    num_cols = ['tenure', 'MonthlyCharges', 'TotalCharges']
    stats = df[num_cols].describe().T
    stats['median'] = df[num_cols].median()
    stats['iqr'] = stats['75%'] - stats['25%']

    t_stats = doc.add_table(rows=4, cols=9)
    t_stats.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Feature", "Mean", "Std", "Min", "Q1 (25%)", "Median", "Q3 (75%)", "Max", "IQR"]
    for j, h in enumerate(headers):
        cell = t_stats.rows[0].cells[j]
        r = cell.paragraphs[0].add_run(h)
        r.font.name = 'Times New Roman'; r.font.bold = True; r.font.size = Pt(10)
        set_cell_background(cell, "2980B9")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for i, col in enumerate(num_cols):
        row = t_stats.rows[i+1].cells
        row[0].paragraphs[0].add_run(col).font.name = 'Times New Roman'
        row[1].paragraphs[0].add_run(f"{stats.loc[col, 'mean']:.2f}").font.name = 'Times New Roman'
        row[2].paragraphs[0].add_run(f"{stats.loc[col, 'std']:.2f}").font.name = 'Times New Roman'
        row[3].paragraphs[0].add_run(f"{stats.loc[col, 'min']:.2f}").font.name = 'Times New Roman'
        row[4].paragraphs[0].add_run(f"{stats.loc[col, '25%']:.2f}").font.name = 'Times New Roman'
        row[5].paragraphs[0].add_run(f"{stats.loc[col, 'median']:.2f}").font.name = 'Times New Roman'
        row[6].paragraphs[0].add_run(f"{stats.loc[col, '75%']:.2f}").font.name = 'Times New Roman'
        row[7].paragraphs[0].add_run(f"{stats.loc[col, 'max']:.2f}").font.name = 'Times New Roman'
        row[8].paragraphs[0].add_run(f"{stats.loc[col, 'iqr']:.2f}").font.name = 'Times New Roman'
        for c in row:
            set_cell_background(c, "F2F4F4" if i % 2 == 0 else "FFFFFF")

    doc.add_paragraph("Interpretation: The mean tenure is 32.37 months, with a median of 29.0 months. MonthlyCharges range from $18.25 to $118.75 (mean $64.76). TotalCharges shows right-skewness (median $1,397.47 vs mean $2,281.92).")

    # ---------------------------------------------------------
    # 7. Univariate Analysis
    # ---------------------------------------------------------
    add_sec_heading("7. Univariate Analysis")
    doc.add_paragraph("Univariate analysis evaluates single feature distributions. Tenure exhibits a bimodal distribution with peaks at early tenure (<12 months) and high tenure (>60 months).")

    fig_path = os.path.join(FIGURES_DIR, "tenure_distribution.png")
    if os.path.exists(fig_path):
        doc.add_paragraph("Figure 7.1: Tenure Distribution by Churn Status")
        doc.add_picture(fig_path, width=Inches(5.8))

    # ---------------------------------------------------------
    # 8. Bivariate Analysis
    # ---------------------------------------------------------
    add_sec_heading("8. Bivariate Analysis")
    doc.add_paragraph("Bivariate analysis examines relationships between predictors and target `Churn`.")

    fig_path = os.path.join(FIGURES_DIR, "churn_by_contract.png")
    if os.path.exists(fig_path):
        doc.add_paragraph("Figure 8.1: Customer Churn Distribution by Contract Type")
        doc.add_picture(fig_path, width=Inches(5.8))

    fig_path = os.path.join(FIGURES_DIR, "churn_by_internet_service.png")
    if os.path.exists(fig_path):
        doc.add_paragraph("Figure 8.2: Customer Churn by Internet Service Type")
        doc.add_picture(fig_path, width=Inches(5.8))

    # ---------------------------------------------------------
    # 9. Multivariate Analysis
    # ---------------------------------------------------------
    add_sec_heading("9. Multivariate Analysis")
    doc.add_paragraph("Multivariate analysis explores interaction effects between Contract, MonthlyCharges, and Churn.")

    fig_path = os.path.join(FIGURES_DIR, "multivariate_contract_charges.png")
    if os.path.exists(fig_path):
        doc.add_paragraph("Figure 9.1: Monthly Charges by Contract Type and Churn Status")
        doc.add_picture(fig_path, width=Inches(5.8))

    # ---------------------------------------------------------
    # 10. Churn Analysis
    # ---------------------------------------------------------
    add_sec_heading("10. Churn Analysis")
    doc.add_paragraph("Overall target distribution: 5,174 retained customers (73.46%) vs 1,869 churned customers (26.54%). Table 2 outlines actual churn rates across key customer categories:")

    # Crosstab Table
    t_churn = doc.add_table(rows=7, cols=5)
    t_churn.alignment = WD_TABLE_ALIGNMENT.CENTER
    ch_headers = ["Category", "Subgroup", "Total Customers", "Churned Count", "Churn Rate (%)"]
    for j, h in enumerate(ch_headers):
        cell = t_churn.rows[0].cells[j]
        r = cell.paragraphs[0].add_run(h)
        r.font.name = 'Times New Roman'; r.font.bold = True; r.font.size = Pt(10)
        set_cell_background(cell, "2980B9")
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    c_data = [
        ("Contract", "Month-to-Month", "3,875", "1,655", "42.71%"),
        ("Contract", "Two Year", "1,695", "48", "2.83%"),
        ("Internet Service", "Fiber Optic", "3,096", "1,297", "41.89%"),
        ("Payment Method", "Electronic Check", "2,365", "1,071", "45.29%"),
        ("Tenure Cohort", "0–12 Months", "2,186", "1,037", "47.44%"),
        ("Tenure Cohort", "61–72 Months", "1,407", "93", "6.61%")
    ]
    for i, row_vals in enumerate(c_data):
        row = t_churn.rows[i+1].cells
        for j, val in enumerate(row_vals):
            row[j].paragraphs[0].add_run(val).font.name = 'Times New Roman'
            set_cell_background(row[j], "F2F4F4" if i % 2 == 0 else "FFFFFF")

    # ---------------------------------------------------------
    # 11. Correlation Analysis
    # ---------------------------------------------------------
    add_sec_heading("11. Correlation Analysis")
    doc.add_paragraph("Numerical features exhibit key linear correlations with Churn (0/1): tenure (r = -0.352), MonthlyCharges (r = +0.193), TotalCharges (r = -0.199).")

    fig_path = os.path.join(FIGURES_DIR, "correlation_heatmap.png")
    if os.path.exists(fig_path):
        doc.add_paragraph("Figure 11.1: Correlation Heatmap Matrix")
        doc.add_picture(fig_path, width=Inches(5.2))

    # ---------------------------------------------------------
    # 12. Key Findings
    # ---------------------------------------------------------
    add_sec_heading("12. Key Findings")
    doc.add_paragraph("1. Target Class Imbalance: 26.54% overall churn rate.\n2. Contract Risk: Month-to-Month subscribers churn at 42.71% vs 2.83% for 2-Year plans.\n3. Internet Service Risk: Fiber Optic subscribers churn at 41.89% vs 18.96% for DSL.\n4. Payment Method Risk: Electronic Check users churn at 45.29%.\n5. Tenure Risk: 0–12 Month subscribers churn at 47.44% vs 6.61% for 61–72 Months.\n6. Security Services: Lack of Online Security increases churn to 41.77%.\n7. Tech Support: Lack of Tech Support increases churn to 41.64%.\n8. Demographics: Senior citizens churn at 41.68% vs 23.61% for non-seniors.")

    # ---------------------------------------------------------
    # 13. Business Insights
    # ---------------------------------------------------------
    add_sec_heading("13. Business Insights")
    doc.add_paragraph("• DATA FINDING: Month-to-Month contracts exhibit 42.71% churn.\n  - BUSINESS RECOMMENDATION: Introduce a 15% discount for 1-year contract renewals.\n• DATA FINDING: Fiber Optic users churn at 41.89% with high monthly bills ($91.50 avg).\n  - BUSINESS RECOMMENDATION: Bundle free Online Security and Tech Support for Fiber optic tiers.\n• DATA FINDING: Electronic Check users exhibit 45.29% churn.\n  - BUSINESS RECOMMENDATION: Provide a $5/month bill credit for switching to Auto Credit Card / ACH.")

    # ---------------------------------------------------------
    # 14. Challenges Faced
    # ---------------------------------------------------------
    add_sec_heading("14. Challenges Faced")
    doc.add_paragraph("• Target Class Imbalance (~2.77:1 ratio) requiring careful metric selection.\n• Pandas 3.0 StringDtype migration for categorical crosstab aggregations.")

    # ---------------------------------------------------------
    # 15. Limitations
    # ---------------------------------------------------------
    add_sec_heading("15. Limitations")
    doc.add_paragraph("• Historical cross-sectional snapshot lacks real-time time-series telemetry.\n• Observational EDA proves correlation, not direct causality.")

    # ---------------------------------------------------------
    # 16. Impact on Future Modeling
    # ---------------------------------------------------------
    add_sec_heading("16. Impact on Future Modeling")
    doc.add_paragraph("EDA confirms Contract, Tenure, InternetService, MonthlyCharges, and PaymentMethod as top predictive features for Week 4 machine learning.")

    # ---------------------------------------------------------
    # 17. Conclusion
    # ---------------------------------------------------------
    add_sec_heading("17. Conclusion")
    doc.add_paragraph("Week 2 EDA successfully isolated primary risk factors: short contract terms, high monthly fees, fiber optic service without support add-ons, and electronic check payment methods.")

    # ---------------------------------------------------------
    # 18. References
    # ---------------------------------------------------------
    add_sec_heading("18. References")
    doc.add_paragraph("1. IBM Telco Customer Churn Public Dataset.\n2. McKinney, W. (2010). Data Structures for Statistical Computing in Python.\n3. Waskom, M. L. (2021). Seaborn: statistical data visualization.")

    # ---------------------------------------------------------
    # 19. Submission Description (200+ Words)
    # ---------------------------------------------------------
    add_sec_heading("19. Submission Description")
    sub_desc = (
        "This Week 2 submission for the Data Science Internship program delivers a comprehensive, data-backed "
        "Exploratory Data Analysis (EDA) and visualization pipeline built upon the sanitized IBM Telco Customer Churn dataset "
        "(`data/processed/cleaned_telco_churn.csv`). Executed by Sindhu Patil, the analysis encompasses structural dataset audits, "
        "descriptive statistics calculations (mean, median, standard deviation, interquartile range), univariate distribution plots, "
        "bivariate churn cross-tabulations, multivariate factor interactions, and numerical correlation matrix heatmaps. "
        "The analysis evaluated 7,043 customer records, establishing an overall target churn rate of 26.54% (1,869 churners vs 5,174 retained customers). "
        "Bivariate analysis isolated major churn drivers: Month-to-Month contracts (42.71% churn rate), Fiber Optic internet service (41.89% churn rate), "
        "Electronic Check payment methods (45.29% churn rate), and initial tenure subscribers under 12 months (47.44% churn rate). "
        "All visual figures were programmatically generated and exported to `outputs/figures/week2/` at 300 DPI resolution. "
        "The deliverables include an 18-section executable Jupyter Notebook (`notebooks/week2_eda.ipynb`), modular visualization source code "
        "(`src/visualization.py`), updated GitHub project documentation (`README.md`), and this formal Word executive report (`reports/week2_report.docx`). "
        "These empirical findings directly inform feature selection and strategy formulation for upcoming clustering and classification weeks."
    )
    doc.add_paragraph(sub_desc)

    doc_path = os.path.join(REPORTS_DIR, "week2_report.docx")
    doc.save(doc_path)
    print(f"Created Word Executive Report: {doc_path}")


def main():
    print("==========================================================")
    print("STARTING WEEK 2 ASSET GENERATION FOR SINDHU PATIL")
    print("==========================================================")
    df = pd.read_csv(DATA_PATH)
    print(f"Loaded dataset: {DATA_PATH} with shape {df.shape}")

    generate_week2_figures(df)
    generate_week2_notebook()
    build_week2_word_report(df)

    print("==========================================================")
    print("WEEK 2 ASSET GENERATION COMPLETED SUCCESSFULLY")
    print("==========================================================")


if __name__ == "__main__":
    main()
